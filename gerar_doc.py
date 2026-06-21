# -*- coding: utf-8 -*-
"""Gera o .docx formatado (ABNT) com as seções atualizadas do TCC I."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf.space_after = Pt(6)


def h(texto, nivel=1):
    doc.add_heading(texto, level=nivel)


def par(texto):
    doc.add_paragraph(texto)


doc.add_heading("TCC I - Secoes atualizadas (sistema Econom-IA)", level=0)
par("Documento de apoio: secoes do TCC I reescritas para refletir o estado atual "
    "do sistema. As secoes 1 (Introducao), 2 (Revisao da Literatura), 4 (Cronograma) "
    "e as Referencias permanecem sem alteracao.")

h("3.5 Alertas integrados", 1)
par("A interface conta com um sistema de alertas visuais. Quando o modelo identifica "
    "um sinal de compra ou venda em uma acao, o usuario visualiza um banner destacado "
    "no Painel Avancado, contendo o nome do ativo, o tipo de sinal, o preco atual e a "
    "estimativa de variacao. A consolidacao desses sinais tambem e apresentada na tela "
    "inicial, em um resumo dos quatro ativos monitorados.")
par("Uma secao dedicada ao historico de alertas, que permitira ao usuario acompanhar "
    "todos os sinais emitidos nos ultimos periodos e realizar a analise retrospectiva "
    "do desempenho do modelo, encontra-se em desenvolvimento, com entrega prevista "
    "para o TC II.")

h("5 RESULTADOS PARCIAIS", 1)
par("Esta secao apresenta os resultados parciais obtidos durante a execucao do TC I, "
    "abrangendo a implementacao da interface web, a estruturacao do pipeline de dados, "
    "o treinamento inicial do modelo de classificacao Random Forest e a analise das "
    "metricas de desempenho geradas. Os resultados aqui descritos validam o "
    "funcionamento integrado de toda a cadeia tecnica concebida e fornecem a base para "
    "os ajustes e comparacoes previstos para a etapa subsequente do trabalho.")

h("5.1 Implementacao da interface", 2)
par("A interface web, denominada Econom-IA, encontra-se funcional no estado atual do "
    "projeto. Adota tema escuro como padrao, decisao tomada com base na reducao do "
    "cansaco visual em sessoes prolongadas de analise e no alinhamento estetico com "
    "plataformas consolidadas do setor financeiro, como Binance, TradingView e o "
    "terminal Bloomberg. Nesta etapa (TC I), o visual foi mantido propositalmente "
    "simples e limpo, reservando o refinamento estetico - identidade visual, logotipo "
    "e aprimoramento do front-end - para o TC II.")
par("A navegacao e organizada em uma barra lateral fixa com tres secoes: Inicio, "
    "Analise e Ferramentas. A secao Inicio concentra a tela de boas-vindas, que combina "
    "uma saudacao contextualizada (variando conforme o turno do dia e o humor do "
    "mercado) e quatro cards com indicadores macroeconomicos atualizados em tempo real: "
    "o IBOVESPA e o dolar (USD/BRL), obtidos por meio da biblioteca yfinance (Yahoo "
    "Finance), e a Taxa Selic e o IPCA acumulado em doze meses, obtidos pela API publica "
    "do Banco Central do Brasil. A tela apresenta ainda um resumo consolidado dos sinais "
    "gerados pelo modelo para os quatro ativos monitorados. Uma area destinada as "
    "manchetes do mercado financeiro esta prevista para o TC II.")
par("Para atender ao publico-alvo heterogeneo do sistema, foram desenvolvidas duas "
    "telas distintas de analise: o Painel Iniciante, voltado a usuarios com pouca "
    "familiaridade no dominio, e o Painel Avancado, destinado a investidores "
    "experientes. O Painel Iniciante apresenta a analise de cada acao com linguagem "
    "cotidiana, perguntas diretas e respostas em texto narrativo, exibindo apenas tres "
    "cartoes principais (preco atual, variacao do ultimo mes e estimativa para os "
    "proximos pregoes). O Painel Avancado exibe todos os indicadores tecnicos com seus "
    "valores numericos brutos, o grafico interativo com os sinais sobrepostos e os "
    "cenarios estatisticos de predicao. Em ambas as telas o usuario pode selecionar o "
    "ativo e o periodo exibido no grafico (ultimo mes, seis meses, um ano ou todo o "
    "historico desde 2019), sendo os ultimos seis meses o periodo exibido por padrao. "
    "Essa abordagem segue o conceito de progressive disclosure no design de interfaces, "
    "no qual a complexidade e revelada gradualmente conforme a necessidade do usuario.")
par("A secao Ferramentas reune recursos complementares - o assistente conversacional "
    "EconomIA, o painel Previsoes & Macro e o Centro de Alertas -, cujas funcionalidades "
    "estao em desenvolvimento, com entrega prevista para o TC II (ver secoes 3.5 e 5.5).")
par("Para garantir desempenho fluido na navegacao, o framework Streamlit utiliza "
    "mecanismos nativos de cache: os dados sao armazenados em memoria apos a primeira "
    "leitura por meio do decorador @st.cache_data, e o modelo treinado e mantido "
    "carregado durante toda a sessao com @st.cache_resource.")

h("5.2 Pipeline de dados e estrutura de armazenamento", 2)
par("Foi realizada a coleta dos dados para os quatro ativos selecionados (ITUB4, BBDC4, "
    "BBAS3 e SANB11), abrangendo o periodo compreendido entre janeiro de 2019 e setembro "
    "de 2025, totalizando 6.968 registros diarios de OHLCV (abertura, maxima, minima, "
    "fechamento e volume). A engenharia de atributos foi aplicada sobre esse conjunto, "
    "produzindo as variaveis derivadas necessarias ao treinamento: RSI de 14 periodos, "
    "MACD com configuracao padrao (12, 26, 9), medias moveis simples e exponenciais de "
    "9, 21 e 50 periodos, e Bandas de Bollinger de 20 periodos. Apos o warm-up dos "
    "indicadores e o tratamento de valores ausentes, o dataset final consolidou-se em "
    "6.732 registros validos com a variavel-alvo de classificacao atribuida.")
par("A arquitetura de armazenamento adotada baseia-se em arquivos persistidos em disco, "
    "sem necessidade de um sistema gerenciador de banco de dados relacional. Os datasets "
    "sao salvos em formato Parquet, que oferece compressao colunar nativa, leitura rapida "
    "em pandas e portabilidade entre maquinas (McKinney, 2010). Apos o treinamento "
    "realizado nesta primeira etapa, o modelo Random Forest e serializado por meio da "
    "biblioteca joblib, gerando o arquivo random_forest.pkl, enquanto as metricas obtidas "
    "sao salvas em arquivo JSON estruturado. A interface consome esses arquivos "
    "diretamente, de modo que qualquer novo treinamento do modelo se reflete "
    "automaticamente nas informacoes apresentadas, sem necessidade de alteracoes no codigo.")

h("5.3 Treinamento do modelo e metricas obtidas", 2)
par("Conforme descrito na secao 3.3, os algoritmos previstos para o trabalho sao Random "
    "Forest, SVM e LSTM. Nesta primeira etapa, foi realizado o treinamento do Random "
    "Forest como ponto de partida para validar a cadeia tecnica completa. O modelo foi "
    "configurado com 150 arvores, profundidade maxima de 10 niveis, min_samples_leaf "
    "igual a 10 e balanceamento automatico das classes para mitigar o desbalanceamento "
    "natural do problema. A acao SANB11 foi reservada integralmente como conjunto de "
    "validacao out-of-sample, totalizando 1.683 registros nunca vistos pelo modelo "
    "durante o aprendizado. Os 5.049 registros restantes, referentes as acoes ITUB4, "
    "BBDC4 e BBAS3, foram divididos em 70% para treino (3.534 registros) e 30% para "
    "teste (1.515 registros), respeitando rigorosamente a ordem cronologica.")
par("Os sinais de compra, venda e neutralidade exibidos na interface sao gerados por "
    "esse modelo treinado, a partir dos indicadores tecnicos calculados para cada ativo, "
    "em conformidade com a metodologia descrita na secao 3.")
par("A validacao cruzada temporal com cinco folds (TimeSeriesSplit) resultou em F1-Score "
    "macro medio de 0,3307, com desvio-padrao de 0,0299 entre os folds. No conjunto de "
    "teste, formado pelas mesmas acoes do treinamento, mas em periodo cronologicamente "
    "posterior, o modelo obteve F1-Score macro de 0,3238, Precision macro de 0,3263, "
    "Recall macro de 0,3234 e acuracia de 0,3254. No conjunto out-of-sample referente a "
    "SANB11, o desempenho foi ligeiramente superior: F1-Score macro de 0,3596, Precision "
    "macro de 0,3601, Recall macro de 0,3600 e acuracia de 0,3678.")

tab = doc.add_table(rows=5, cols=3)
tab.style = "Light Grid Accent 1"
tab.alignment = WD_TABLE_ALIGNMENT.CENTER
dados = [
    ("Metrica (macro)", "Conjunto de teste", "Out-of-sample (SANB11)"),
    ("F1-Score", "0,3238", "0,3596"),
    ("Precision", "0,3263", "0,3601"),
    ("Recall", "0,3234", "0,3600"),
    ("Acuracia", "0,3254", "0,3678"),
]
for i, linha in enumerate(dados):
    for j, valor in enumerate(linha):
        cel = tab.rows[i].cells[j]
        cel.text = valor
        for p in cel.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
                if i == 0:
                    r.font.bold = True
doc.add_paragraph("")

par("A analise da importancia das features indicou que os tres indicadores tecnicos com "
    "maior contribuicao nas decisoes do modelo foram o z-score do histograma do MACD "
    "(15,99%), a razao entre o preco de fechamento e a media movel simples de 9 periodos "
    "(13,35%) e a razao entre o preco de fechamento e a media movel simples de 50 "
    "periodos (12,87%); o RSI de 14 periodos aparece logo em seguida (12,14%). Em "
    "conjunto, os tres principais indicadores respondem por aproximadamente 42% da "
    "capacidade discriminativa do modelo, validando empiricamente a relevancia da "
    "combinacao entre indicadores de momento (MACD) e de tendencia (medias moveis), ja "
    "estabelecida na literatura de analise tecnica (Achelis, 2013).")

h("5.4 Analise dos resultados parciais", 2)
par("Os valores obtidos nesta primeira execucao do treinamento encontram-se proximos do "
    "desempenho esperado de um classificador aleatorio, que seria de aproximadamente "
    "0,333 para um problema balanceado de tres classes. Esse resultado, embora modesto "
    "em termos absolutos, e considerado coerente e satisfatorio para esta fase do "
    "projeto, pois confirma o funcionamento integrado de toda a cadeia tecnica concebida: "
    "a coleta dos dados via API, o calculo correto dos indicadores tecnicos, a criacao "
    "adequada da variavel-alvo, a divisao temporal respeitada, o treinamento estavel do "
    "modelo, a persistencia dos artefatos e o consumo automatizado das metricas pela "
    "interface.")
par("Observou-se ainda que o desempenho no conjunto out-of-sample (SANB11) foi "
    "ligeiramente superior ao desempenho no conjunto de teste com as mesmas acoes do "
    "treinamento, o que sugere ausencia de overfitting significativo aos ativos "
    "especificos utilizados durante o aprendizado. Esse comportamento corrobora a "
    "estrategia de treinamento generalista descrita por Najem et al. (2024), segundo a "
    "qual modelos treinados com dados de multiplos ativos tendem a aprender padroes mais "
    "universais dos indicadores tecnicos.")
par("Considerando o desempenho ainda proximo do aleatorio, fica evidente a necessidade "
    "de ajustes adicionais nas proximas etapas do projeto. As principais direcoes "
    "identificadas para a continuidade do trabalho incluem o ajuste sistematico de "
    "hiperparametros por meio de busca em grade ou validacao bayesiana, a inclusao de "
    "novas variaveis explicativas (volume normalizado, volatilidade e defasagens do "
    "retorno), a comparacao experimental com os modelos SVM e LSTM previstos no escopo "
    "original, e a reavaliacao do limiar de classificacao atualmente fixado em +/-1,5%.")

h("5.5 Painel macroeconomico e integracao com mercado de previsao", 2)
par("Foi planejada uma tela complementar denominada Previsoes & Macro, voltada a "
    "contextualizacao macroeconomica dos sinais gerados pelo modelo. Essa tela reunira "
    "cards informativos com as principais taxas e indicadores brasileiros relevantes "
    "para o setor bancario (Taxa Selic, IPCA acumulado em doze meses, Taxa do Fed Funds "
    "americano e cotacao do dolar comercial); cards com os rendimentos dos titulos do "
    "Tesouro americano (Treasury Yields); um grafico historico comparativo entre a Taxa "
    "Selic e o IPCA; e um painel dedicado ao mercado de previsao Polymarket, exibindo as "
    "probabilidades agregadas para a proxima decisao do Copom sobre a Taxa Selic.")
par("A incorporacao do Polymarket representa uma das contribuicoes originais deste "
    "trabalho. Mercados de previsao (prediction markets) sao instrumentos em que "
    "participantes apostam valores reais sobre o resultado de eventos futuros, e cujos "
    "precos agregados refletem com notavel precisao a expectativa coletiva dos agentes "
    "(Wolfers; Zitzewitz, 2004).")
par("Cabe destacar que esses indicadores macroeconomicos tem carater exclusivamente "
    "contextual e nao sao utilizados como entrada do modelo de classificacao - os sinais "
    "de compra e venda derivam unicamente dos indicadores tecnicos processados pelo "
    "Random Forest. Esta tela encontra-se em desenvolvimento: a automatizacao da leitura "
    "dos indicadores por meio das APIs do Banco Central do Brasil, do sistema FRED "
    "(Federal Reserve Economic Data) e do Polymarket esta prevista para o TC II.")

doc.save("TCC_I_Econom-IA_atualizado.docx")
print("Arquivo .docx gerado com sucesso.")